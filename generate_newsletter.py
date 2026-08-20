"""
Builds the Catholic Newsletter (Blessed Sacrament Catholic Church) as a Navy & Gold .docx file.

Usage:
    python generate_newsletter.py content.json

Reads a JSON file describing the week's newsletter content (see content.example.json
for the expected shape) and writes a formatted .docx into the issues/ folder.
"""

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x00, 0x00, 0x80)
GOLD = RGBColor(0xFF, 0xD7, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
BODY_TEXT = RGBColor(0x22, 0x22, 0x22)

NAVY_HEX = "000080"
GOLD_HEX = "FFD700"

FONT_HEADING = "Georgia"
FONT_BODY = "Calibri"

# FFD700 (bright gold) reads well on the navy banner but is too pale for body
# text on a white page, so meta lines (source | date) use a darker goldenrod
# for legibility while borders/rules still use the true banner gold.
GOLD_TEXT = RGBColor(0xB8, 0x86, 0x0B)


def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_border(paragraph, hex_color, side="bottom", size=12, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), hex_color)
    pBdr.append(border)
    pPr.append(pBdr)


def set_narrow_margins(doc, inches=0.9):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


SPANISH_MONTHS = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


def format_date_line(content):
    raw = content.get("date_produced")
    if raw:
        return raw
    issue_date = content.get("issue_date")
    if not issue_date:
        return ""
    try:
        dt = datetime.strptime(issue_date, "%Y-%m-%d")
    except ValueError:
        return issue_date
    if content.get("language", "en").split("-")[0].lower() == "es":
        return f"{dt.day} de {SPANISH_MONTHS[dt.month]} de {dt.year}"
    return dt.strftime("%B %d, %Y")


def add_masthead(doc, title, parish_line, date_line):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, NAVY_HEX)
    cell.width = Inches(6.5)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.font.color.rgb = GOLD
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = FONT_HEADING

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(parish_line)
    run2.font.color.rgb = WHITE
    run2.font.size = Pt(13)
    run2.italic = True
    run2.font.name = FONT_HEADING

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    run3 = p3.add_run(date_line)
    run3.font.color.rgb = GOLD
    run3.font.size = Pt(11)
    run3.font.name = FONT_HEADING

    accent = doc.add_paragraph()
    accent.paragraph_format.space_after = Pt(16)
    add_border(accent, GOLD_HEX, side="bottom", size=14, space=2)


def add_intro(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = FONT_BODY


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.color.rgb = NAVY
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.name = FONT_HEADING
    add_border(p, GOLD_HEX, side="bottom", size=10, space=4)


def add_article(doc, article):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(1)
    run_title = p_title.add_run(article.get("title", "").strip())
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = NAVY
    run_title.font.name = FONT_BODY

    meta_bits = []
    if article.get("source"):
        meta_bits.append(article["source"])
    if article.get("date"):
        meta_bits.append(article["date"])
    if meta_bits:
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(3)
        run_meta = p_meta.add_run(" | ".join(meta_bits))
        run_meta.italic = True
        run_meta.font.size = Pt(9)
        run_meta.font.color.rgb = GOLD_TEXT
        run_meta.font.name = FONT_BODY

    if article.get("summary"):
        p_summary = doc.add_paragraph()
        p_summary.paragraph_format.space_after = Pt(4)
        run_summary = p_summary.add_run(article["summary"].strip())
        run_summary.font.size = Pt(10.5)
        run_summary.font.color.rgb = BODY_TEXT
        run_summary.font.name = FONT_BODY

    if article.get("url"):
        p_link = doc.add_paragraph()
        p_link.paragraph_format.space_after = Pt(2)
        run_link = p_link.add_run(article["url"])
        run_link.font.size = Pt(9)
        run_link.font.color.rgb = RGBColor(0x1A, 0x4B, 0x8C)
        run_link.font.name = FONT_BODY
        run_link.underline = True


def add_footer_note(doc, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    add_border(p, GOLD_HEX, side="top", size=10, space=8)
    run = p.add_run(note)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    run.font.name = FONT_BODY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build(content, out_path):
    doc = Document()
    set_narrow_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)

    add_masthead(
        doc,
        content.get("title", "Catholic Newsletter"),
        content.get("subtitle", "Blessed Sacrament Catholic Church"),
        format_date_line(content),
    )
    add_intro(doc, content.get("intro"))

    for section in content.get("sections", []):
        add_section_heading(doc, section.get("heading", ""))
        for article in section.get("articles", []):
            add_article(doc, article)

    add_footer_note(
        doc,
        content.get(
            "footer_note",
            "Compiled from Vatican and Catholic press sources. For personal use.",
        ),
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    if len(sys.argv) != 2:
        print("Usage: python generate_newsletter.py <content.json>")
        sys.exit(1)

    content_path = Path(sys.argv[1])
    content = json.loads(content_path.read_text(encoding="utf-8"))

    issue_date = content.get("issue_date", "issue")
    language = content.get("language", "en").split("-")[0].lower()
    suffix = "" if language == "en" else f"_{language.upper()}"
    out_path = Path("issues") / f"Catholic_Intelligence_Newsletter_{issue_date}{suffix}.docx"

    saved_path = build(content, out_path)
    print(f"Saved: {saved_path.resolve()}")


if __name__ == "__main__":
    main()
